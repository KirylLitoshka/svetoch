from aiohttp import web
from xml.etree import ElementTree as ET
from utils import MONTHS
from docx import Document
from docx.shared import Inches, Pt
from zipfile import ZipFile

import aiofiles
import calendar
import os
from datetime import date


async def build_consolidated_report(report_data, month, year, currency_coefficient):
    """
    Построитель сводной ведомости
    """
    month_name = MONTHS[month - 1]
    file_path = "warmth/reports/templates/consolidated_report.txt"
    output_file_path = "warmth/reports/out/consolidated_report.txt"
    title_line = "│{:3d}│{:26s}│{:12.4f}│{:14.2f}│{:12.4f}│{:14.2f}│{:12.4f}│{:14.2f}│"
    sub_line = "│{:3s}│ - {:23s}│{:12.4f}│{:14.2f}│{:12.4f}│{:14.2f}│{:12.4f}│{:14.2f}│"
    empty_line = "│{:3s}│{:26s}│{:12s}│{:14s}│{:12s}│{:14s}│{:12s}│{:14s}│"
    end_line = "│{:3s}│{:26s}│{:12.4f}│{:14.2f}│{:12.4f}│{:14.2f}│{:12.4f}│{:14.2f}│"
    total = {
        "heating_value": sum([item["heating_value"] for item in report_data]),
        "heating_cost": sum([item["heating_cost"] for item in report_data]),
        "water_heating_value": sum(
            [item["water_heating_value"] for item in report_data]
        ),
        "water_heating_cost": sum([item["water_heating_cost"] for item in report_data]),
        "total_value": sum([item["total_value"] for item in report_data]),
        "total_cost": sum([item["total_cost"] for item in report_data]),
    }
    main_content = ""
    for index, item in enumerate(report_data):
        entries = item.pop("workshops")
        main_content += title_line.format(index + 1, *item.values()) + "\n"
        for entry in entries:
            for key in ["id", "workshop_group_id", "is_currency_coefficient_applied"]:
                entry.pop(key)
            main_content += sub_line.format("", *entry.values()) + "\n"
        main_content += empty_line.format(*[""] * 8)
        if index != len(report_data) - 1:
            main_content += "\n"
    end_content = end_line.format("", "Итого", *total.values())

    async with aiofiles.open(file_path, mode="r", encoding="utf8") as fp:
        template_data = await fp.read()
    data = template_data.format(
        month_name, year, currency_coefficient["value_1"], main_content, end_content
    )
    async with aiofiles.open(output_file_path, mode="w", encoding="utf8") as output_fp:
        await output_fp.write(data)
    return web.FileResponse(
        path=output_file_path,
        status=200,
        headers={
            "Content-Disposition": f"attachment;filename=consolidated_report_{month}_{year}.txt"
        },
    )


async def build_workshop_report(data, month, year):
    month_name = MONTHS[month - 1]
    file_path = f"warmth/reports/templates/workshop_report.txt"
    output_file_path = f"warmth/reports/out/workshop_report.txt"
    line = "│{:27s}│{:11.4f}│{:11.2f}│{:11.4f}│{:11.2f}│{:11.4f}│{:11.2f}│"
    main_content = ""
    for index, row in enumerate(data["objects"]):
        main_content += line.format(*row.values())
        main_content += "" if index == len(data["objects"]) - 1 else "\n"
    end_line = line.format("Итого", *list(data.values())[3:])
    async with aiofiles.open(file_path, mode="r", encoding="utf8") as fp:
        template_data = await fp.read()
    data = template_data.format(data["title"], month_name, year, main_content, end_line)
    async with aiofiles.open(output_file_path, mode="w", encoding="utf8") as output_fp:
        await output_fp.write(data)
    return web.FileResponse(
        output_file_path,
        status=200,
        headers={
            "Content-Disposition": f"attachment;filename=workshop_report_{month}_{year}.txt"
        },
    )


async def build_renter_short_report(report_data, month, year):
    month_name = MONTHS[month - 1]
    file_path = f"warmth/reports/templates/renter_short_report.txt"
    output_file_path = f"warmth/reports/out/renter_short_report.txt"
    line = "│{:3s}│{:18s}│{:3s}│{:13.4f}│{:13.2f}│{:13.2f}│{:13.2f}│{:13.2f}│"
    end_line = "│{:22s}│{:3s}│{:13.4f}│{:13.2f}│{:13.2f}│{:13.2f}│{:13.2f}│"
    content, end_content = "", ""

    for index, row in enumerate(report_data):
        content += (
            line.format(
                str(index + 1),
                row["title"][:18],
                "",
                round(row["heating"]["value"] + row["water_heating"]["value"], 4),
                round(row["heating"]["cost"] + row["water_heating"]["cost"], 2),
                round(
                    row["heating"]["coefficient"] + row["water_heating"]["coefficient"],
                    2,
                ),
                round(row["heating"]["vat"] + row["water_heating"]["vat"], 2),
                round(row["heating"]["total"] + row["water_heating"]["total"], 2),
            )
            + "\n"
        )
        if row["heating"]["total"] or row["heating"]["cost"]:
            content += (
                line.format(
                    "",
                    "",
                    "Отп",
                    round(row["heating"]["value"], 4),
                    round(row["heating"]["cost"], 2),
                    round(row["heating"]["coefficient"], 2),
                    round(row["heating"]["vat"], 2),
                    round(row["heating"]["total"], 2),
                )
                + "\n"
            )
        if row["water_heating"]["total"] or row["water_heating"]["cost"]:
            content += (
                line.format(
                    "",
                    "",
                    "ГВС",
                    round(row["water_heating"]["value"], 4),
                    round(row["water_heating"]["cost"], 2),
                    round(row["water_heating"]["coefficient"], 2),
                    round(row["water_heating"]["vat"], 2),
                    round(row["water_heating"]["total"], 2),
                )
                + "\n"
            )
    content += "│{:3s}│{:18s}│{:3s}│{:13s}│{:13s}│{:13s}│{:13s}│{:13s}│".format(
        "", "", "", "", "", "", "", ""
    )

    end_content += (
        end_line.format(
            "Всего начислено",
            "",
            round(
                sum(
                    [
                        row["heating"]["value"] + row["water_heating"]["value"]
                        for row in report_data
                    ]
                ),
                4,
            ),
            round(
                sum(
                    [
                        row["heating"]["cost"] + row["water_heating"]["cost"]
                        for row in report_data
                    ]
                ),
                2,
            ),
            round(
                sum(
                    [
                        row["heating"]["coefficient"]
                        + row["water_heating"]["coefficient"]
                        for row in report_data
                    ]
                ),
                2,
            ),
            round(
                sum(
                    [
                        row["heating"]["vat"] + row["water_heating"]["vat"]
                        for row in report_data
                    ]
                ),
                2,
            ),
            round(
                sum(
                    [
                        row["heating"]["total"] + row["water_heating"]["total"]
                        for row in report_data
                    ]
                ),
                2,
            ),
        )
        + "\n"
    )
    end_content += (
        end_line.format(
            "в том числе отопление.",
            "Гкл",
            round(sum([row["heating"]["value"] for row in report_data]), 4),
            round(sum([row["heating"]["cost"] for row in report_data]), 2),
            round(sum([row["heating"]["coefficient"] for row in report_data]), 2),
            round(sum([row["heating"]["vat"] for row in report_data]), 2),
            round(sum([row["heating"]["total"] for row in report_data]), 2),
        )
        + "\n"
    )
    end_content += end_line.format(
        "в том числе подог.воды",
        "Гкл",
        round(sum([row["water_heating"]["value"] for row in report_data]), 4),
        round(sum([row["water_heating"]["cost"] for row in report_data]), 2),
        round(sum([row["water_heating"]["coefficient"] for row in report_data]), 2),
        round(sum([row["water_heating"]["vat"] for row in report_data]), 2),
        round(sum([row["water_heating"]["total"] for row in report_data]), 2),
    )
    async with aiofiles.open(file_path, mode="r", encoding="utf8") as fp:
        template_data = await fp.read()
    data = template_data.format(month_name, year, content, end_content)
    async with aiofiles.open(output_file_path, mode="w", encoding="utf8") as output_fp:
        await output_fp.write(data)
    return web.FileResponse(
        output_file_path,
        status=200,
        headers={
            "Content-Disposition": f"attachment;filename=renter_report_mini_{month}_{year}.txt"
        },
    )


async def build_renter_full_report(renter_payments, month, year):
    month_name = MONTHS[month - 1]
    file_path = f"warmth/reports/templates/renter_full_report.txt"
    output_file_path = f"warmth/reports/out/renter_full_report.txt"
    line = "│{:6s}│{:20s}│{:3s}│{:8s}│{:5s}│{:9s}│{:9s}│{:11s}│{:9s}│{:10s}│{:11s}│"
    sub_line = "│     -│{:20s}│{:3s}│{:8.6f}│{:5.2f}│{:9s}│{:9s}│{:11.2f}│{:9s}│{:10.2f}│{:11.2f}│"
    end_lind = "│------│{:20s}│{:3s}│{:8s}│{:5s}│{:9s}│{:9s}│{:11.2f}│{:9.2f}│{:10.2f}│{:11.2f}│"
    content = ""
    total = {"cost": 0, "vat": 0, "coefficient": 0, "summary": 0}

    for index, renter in enumerate(renter_payments):
        content += line.format(str(index + 1), renter["name"][:20], *[""] * 10) + "\n"
        for payment in renter["payments"]:
            if payment["is_additional_coefficient_applied"]:
                coefficient = payment["additional_coefficient_value"]
            else:
                coefficient = payment["coefficient_value"]
            if payment["heating_cost"]:
                value_with_coefficient = round(
                    round(payment["heating_cost"], 2) * coefficient, 2
                )
                vat = (
                    round(value_with_coefficient / 100 * payment["vat"], 2)
                    if payment["vat"]
                    else 0
                )
                payment_coefficient = round(
                    value_with_coefficient - round(payment["heating_cost"], 2), 2
                )
                total_cost = round(value_with_coefficient + vat, 2)
                total["cost"] += round(payment["heating_cost"], 2)
                total["vat"] += vat
                total["coefficient"] += payment_coefficient
                total["summary"] += total_cost
                content += (
                    sub_line.format(
                        payment["title"][:20],
                        "Отп",
                        coefficient,
                        payment["vat"],
                        str(round(payment["heating_value"], 4)),
                        str(round(payment["applied_rate_value"], 5)),
                        payment["heating_cost"],
                        str(payment_coefficient),
                        vat,
                        total_cost,
                    )
                    + "\n"
                )
            if payment["water_heating_cost"]:
                value_with_coefficient = round(
                    round(payment["water_heating_cost"], 2) * coefficient, 2
                )
                vat = (
                    round(value_with_coefficient / 100 * payment["vat"], 2)
                    if payment["vat"]
                    else 0
                )
                payment_coefficient = round(
                    value_with_coefficient - round(payment["water_heating_cost"], 2), 2
                )
                total_cost = round(value_with_coefficient + vat, 2)
                total["vat"] += vat
                total["cost"] += round(payment["water_heating_cost"], 2)
                total["coefficient"] += payment_coefficient
                total["summary"] += total_cost
                content += (
                    sub_line.format(
                        payment["title"][:20],
                        "ГВС",
                        coefficient,
                        payment["vat"],
                        str(round(payment["water_heating_value"], 4)),
                        str(round(payment["applied_rate_value"], 6)),
                        payment["water_heating_cost"],
                        str(payment_coefficient),
                        vat,
                        total_cost,
                    )
                    + "\n"
                )
    content += line.format(*[""] * 12)
    total_line = end_lind.format(
        "Итого",
        "",
        "",
        "",
        "",
        "",
        total["cost"],
        total["coefficient"],
        total["vat"],
        total["summary"],
    )
    async with aiofiles.open(file_path, mode="r", encoding="utf8") as fp:
        template_data = await fp.read()
    data = template_data.format(month_name, year, content, total_line)
    async with aiofiles.open(output_file_path, mode="w", encoding="utf8") as output_fp:
        await output_fp.write(data)
    return web.FileResponse(
        output_file_path,
        status=200,
        headers={
            "Content-Disposition": f"attachment;filename=renter_report_full_{month}_{year}.txt"
        },
    )


async def build_renter_bank_report(renters_payments, month, year):
    month_name = MONTHS[month - 1]
    last_day_of_month = calendar.monthrange(year, month)[1]
    output_file_path = "warmth/reports/out/renter_bank_invoice.txt"
    invoice_date = "{:2d}.{:02d}.{:4d}".format(last_day_of_month, month, year)
    output_content = ""

    async with aiofiles.open(
        "warmth/reports/templates/renter_bank_invoice.txt", mode="r", encoding="utf-8"
    ) as fp:
        template_data = await fp.read()

    for renter in renters_payments:
        invoice_number = "{:1s}{:02d}{:04d}".format(str(year)[-1], month, renter["id"])
        renter_summary_cost = 0
        renter_summary_currency = 0
        renter_summary_vat = 0
        renter_total = 0
        payments_info = f""

        for index, payment in enumerate(renter["payments"]):
            payments_info += f" {payment['title']}\n"

            if payment["heating_cost"]:
                payments_info += build_renter_bank_report_line(payment, "heating")
                renter_summary_cost += round(payment["heating_cost"], 2)
                renter_summary_currency += payment["heating_coefficient_value"]
                renter_summary_vat += payment["heating_vat_cost"]
                renter_total += payment["heating_total_cost"]

            if payment["water_heating_cost"]:
                payments_info += build_renter_bank_report_line(payment, "water_heating")
                renter_summary_cost += round(payment["water_heating_cost"], 2)
                renter_summary_currency += payment["water_heating_coefficient_value"]
                renter_summary_vat += payment["water_heating_vat_cost"]
                renter_total += payment["water_heating_total_cost"]

            if index != len(renter["payments"]) - 1:
                payments_info += "\n"

        output_content += (
            template_data.format(
                invoice_date=invoice_date,
                invoice_number=invoice_number,
                banking_account=renter["banking_account"],
                registration_number=renter["registration_number"],
                full_name=renter["full_name"],
                contract_number=renter["contract_number"],
                contract_date=renter["contract_date"].strftime("%d.%m.%Y"),
                bank_code=renter["bank_code"],
                bank_title=renter["bank_title"],
                month_name=month_name,
                year=year,
                sum=renter_summary_cost,
                currency=renter_summary_currency,
                vat_sum=renter_summary_vat,
                payment_sum=renter_total,
                payments_detail=payments_info,
            )
            + "\n"
        )

    async with aiofiles.open(
        output_file_path, mode="w", encoding="cp1251", errors="ignore"
    ) as output_fp:
        await output_fp.write(output_content)

    return web.FileResponse(
        output_file_path,
        status=200,
        headers={
            "Content-Disposition": f"attachment;filename=bank_{month}_{year}.txt",
            "Access-Control-Expose-Headers": "Content-Disposition",
        },
    )


async def build_renters_invoices_report(renters, month, year):
    invoice_number = "400004956-{:4d}-5{:2s}{:02d}5{:04d}"
    last_day_of_month = calendar.monthrange(year, month)[1]
    invoice_date = "{:4d}-{:02d}-{:02d}".format(year, month, last_day_of_month)
    output_file_path = "warmth/reports/out/invoices.zip"

    if os.path.exists(output_file_path):
        os.remove(output_file_path)

    zip_file = ZipFile(output_file_path, "a")

    def build_roster_item(
        element, index, title, value, rate_value, service_cost, coefficient, vat
    ):
        roster_item = ET.SubElement(element, "rosterItem")
        payment_cost = round(round(service_cost, 2) * coefficient, 2)
        payment_vat = round(payment_cost / 100 * vat, 2)
        ET.SubElement(roster_item, "number").text = str(index)
        ET.SubElement(roster_item, "name").text = title
        if value:
            ET.SubElement(roster_item, "units").text = "233"
        ET.SubElement(roster_item, "count").text = str(value)
        ET.SubElement(roster_item, "price").text = str(rate_value)
        ET.SubElement(roster_item, "cost").text = str(payment_cost)
        ET.SubElement(roster_item, "summaExcise").text = "0"
        vat_element = ET.SubElement(roster_item, "vat")
        if vat:
            ET.SubElement(vat_element, "rate").text = "{:2.2f}".format(vat)
            ET.SubElement(vat_element, "rateType").text = "DECIMAL"
            ET.SubElement(vat_element, "summaVat").text = str(payment_vat)
        ET.SubElement(roster_item, "costVat").text = str(
            round(payment_cost + payment_vat, 2)
        )

    async with aiofiles.open("warmth/reports/templates/renter_invoice.txt") as f:
        invoice_template = await f.read()

    for renter in renters:
        renter_invoice_number = invoice_number.format(
            year, str(year)[-2:], month, renter["id"]
        )
        renter_short_invoice_number = "{:1s}{:02d}{:04d}".format(
            str(year)[-1], month, renter["id"]
        )
        renter_invoice_path = f"warmth/reports/out/{renter_invoice_number}.xml"

        total_currency_cost = sum(
            [payment["currency_cost"] for payment in renter["payments"]]
        )
        cost = sum(
            [
                round(payment["heating_cost"], 2)
                + round(payment["water_heating_cost"], 2)
                for payment in renter["payments"]
            ]
        )
        total_cost = round(cost + total_currency_cost, 2)
        total_vat = round(
            sum([payment["vat_value"] for payment in renter["payments"]]), 2
        )
        summary = round(total_cost + total_vat, 2)

        invoice_text = invoice_template.format(
            invoice_number=renter_invoice_number,
            invoice_date=invoice_date,
            renter_name=renter["full_name"],
            renter_address=renter["address"],
            contract_date=renter["contract_date"],
            contract_number=renter["contract_number"],
            renter_registration_number=renter["registration_number"] or "",
            short_invoice_number=renter_short_invoice_number,
            total_cost=total_cost,
            total_vat=total_vat,
            summary=summary,
        )

        xml_tree = ET.ElementTree(ET.fromstring(invoice_text))
        xml_tree_root = xml_tree.getroot()
        xml_tree_root.set("xmlns", "http://www.w3schools.com")
        xml_tree_root.set("sender", "400004956")
        roster = xml_tree_root.find("roster")
        roster_index = 1

        for payment in renter["payments"]:
            payment_coefficient = payment["coefficient_value"]
            if payment["is_additional_coefficient_applied"]:
                payment_coefficient = payment["additional_coefficient_value"]

            if payment["heating_cost"]:
                build_roster_item(
                    roster,
                    roster_index,
                    "Отопление",
                    payment["heating_value"],
                    payment["applied_rate_value"],
                    payment["heating_cost"],
                    payment_coefficient,
                    payment["vat"],
                )
                roster_index += 1

            if payment["water_heating_cost"]:
                build_roster_item(
                    roster,
                    roster_index,
                    "Подогр.воды",
                    payment["water_heating_value"],
                    payment["applied_rate_value"],
                    payment["water_heating_cost"],
                    payment_coefficient,
                    payment["vat"],
                )
                roster_index += 1

        ET.indent(xml_tree, space="\t")

        xml_tree.write(renter_invoice_path, encoding="utf-8")
        zip_file.write(renter_invoice_path, f"{renter_invoice_number}.xml")
        os.remove(renter_invoice_path)

    zip_file.close()
    return web.FileResponse(
        output_file_path,
        status=200,
        headers={
            "Content-Disposition": f"attachment;filename=invoices_{month}_{year}.zip",
            "Access-Control-Expose-Headers": "Content-Disposition",
        },
    )


async def build_renters_invoices_print_report(renters_payments, month, year, details):
    month_title = MONTHS[month - 1]
    last_day_of_month = calendar.monthrange(year, month)[1]
    output_file_path = "warmth/reports/out/invoices.docx"
    content_line = "|{:12s}|{:5s}|{:9.5f}|{:12.5f}|{:11.2f}|{:6.2f}|{:7.2f}|{:9.2f}|"
    add_content_line = " {:3d} {:18s} {:7s} {:7s} {:7.6f} {:4.2f} {:9.4f} {:9.5f} {:11.2f} {:10.2f} {:10.2f} {:11.2f}"

    async with aiofiles.open("warmth/reports/templates/renter_invoice_print.txt") as f:
        invoice_text_template = await f.read()

    async with aiofiles.open(
        "warmth/reports/templates/renter_invoice_print_add.txt"
    ) as f:
        additional_invoice_text_template = await f.read()

    document = Document()
    section = document.sections[-1]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.3)
    style = document.styles["Normal"]
    style.font.name = "Arial"

    for index, renter in enumerate(renters_payments):
        renter_short_invoice_number = "{:1s}{:02d}{:04d}".format(
            str(year)[-1], month, renter["id"]
        )
        total = {
            "cost": 0,
            "vat": 0,
            "value": 0,
        }
        heating = total.copy()
        water_heating = total.copy()
        invoice_text_content = ""
        additional_invoice_text_content = ""

        for payment in renter["payments"]:
            tab = 1
            payment_coefficient = payment["coefficient_value"]
            if payment["is_additional_coefficient_applied"]:
                payment_coefficient = payment["additional_coefficient_value"]
            if payment["heating_cost"]:
                value_with_coefficient = (
                    round(payment["heating_cost"], 2) * payment_coefficient
                )
                vat = (
                    round(value_with_coefficient / 100 * payment["vat"])
                    if payment["vat"]
                    else 0
                )
                coefficient_cost = round(
                    value_with_coefficient - round(payment["heating_cost"], 2), 2
                )
                total_cost = round(value_with_coefficient + vat, 2)
                heating["cost"] += value_with_coefficient
                heating["vat"] += vat
                heating["value"] += round(payment["heating_value"], 5)
                additional_invoice_text_content += (
                    add_content_line.format(
                        tab,
                        payment["title"][:18],
                        "Отп/Гкл",
                        "{:02d}.{:4d}".format(month, year),
                        payment_coefficient,
                        payment["vat"],
                        payment["heating_value"],
                        payment["applied_rate_value"],
                        round(payment["heating_cost"], 2),
                        coefficient_cost,
                        vat,
                        total_cost,
                    )
                    + "\n"
                )
                tab += 1
            if payment["water_heating_cost"]:
                value_with_coefficient = (
                    round(payment["water_heating_cost"], 2) * payment_coefficient
                )
                vat = (
                    round(value_with_coefficient / 100 * payment["vat"], 2)
                    if payment["vat"]
                    else 0
                )
                coefficient_cost = round(
                    value_with_coefficient - round(payment["water_heating_cost"], 2), 2
                )
                total_cost = round(value_with_coefficient + vat, 2)
                water_heating["cost"] += value_with_coefficient
                water_heating["vat"] += vat
                water_heating["value"] += round(payment["water_heating_value"], 5)
                additional_invoice_text_content += (
                    add_content_line.format(
                        tab,
                        payment["title"][:18],
                        "ГВС/Гкл",
                        "{:02d}.{:4d}".format(month, year),
                        payment_coefficient,
                        payment["vat"],
                        payment["water_heating_value"],
                        payment["applied_rate_value"],
                        round(payment["water_heating_cost"], 2),
                        coefficient_cost,
                        vat,
                        total_cost,
                    )
                    + "\n"
                )
                tab += 1

        max_rate_value = max([row["applied_rate_value"] for row in renter["payments"]])
        invoice_text_content += (
            content_line.format(
                "Отопление",
                "Гкл",
                max_rate_value,
                heating["value"],
                heating["cost"],
                20,
                heating["vat"],
                round(heating["cost"] + heating["vat"], 2),
            )
            + "\n"
        )
        invoice_text_content += content_line.format(
            "Подогр.воды",
            "Гкл",
            max_rate_value,
            water_heating["value"],
            water_heating["cost"],
            20,
            water_heating["vat"],
            round(water_heating["cost"] + water_heating["vat"], 2),
        )

        renter_invoice_details = next(
            (row for row in details if row["id"] == renter["id"]), None
        )
        if renter_invoice_details is None:
            raise

        if renter_invoice_details["invoice"]:
            document.add_paragraph().add_run(
                invoice_text_template.format(
                    invoice_number=renter_short_invoice_number,
                    day=last_day_of_month,
                    month="{:02d}".format(month),
                    month_title=month_title,
                    year=year,
                    contract_date=renter["contract_date"].strftime("%d.%m.%Y"),
                    contract_number=renter["contract_number"],
                    renter_title=renter["full_name"],
                    renter_address=renter["address"],
                    renter_registration_number=renter["registration_number"] or "",
                    renter_banking_account=renter["banking_account"] or "",
                    bank_code=renter["bank_code"] or "",
                    bank_title=renter["bank_title"] or "",
                    content=invoice_text_content,
                    total_cost=round(heating["cost"] + water_heating["cost"], 2),
                    total_vat=round(heating["vat"] + water_heating["vat"], 2),
                    summary=round(
                        heating["cost"]
                        + heating["vat"]
                        + water_heating["cost"]
                        + water_heating["vat"],
                        2,
                    ),
                ),
                style="Macro Text Char",
            ).font.size = Pt(10)

            document.add_page_break()

        if renter_invoice_details["attachment"]:
            document.add_paragraph().add_run(
                additional_invoice_text_template.format(
                    invoice_number=renter_short_invoice_number,
                    day=last_day_of_month,
                    month="{:02d}".format(month),
                    month_title=month_title,
                    year=year,
                    renter_title=renter["full_name"],
                    renter_address=renter["address"],
                    renter_registration_number=renter["registration_number"] or "",
                    renter_banking_account=renter["banking_account"] or "",
                    bank_code=renter["bank_code"] or "",
                    bank_title=renter["bank_title"] or "",
                    total_vat=round(heating["vat"] + water_heating["vat"], 2),
                    summary=round(
                        heating["cost"]
                        + heating["vat"]
                        + water_heating["cost"]
                        + water_heating["vat"],
                        2,
                    ),
                    content=additional_invoice_text_content,
                ),
                style="Macro Text Char",
            ).font.size = Pt(7)

            if index != len(renters_payments) - 1:
                document.add_page_break()

    document.save(output_file_path)
    return web.FileResponse(
        output_file_path,
        status=200,
        headers={
            "Content-Disposition": f"attachment;filename=invoices_{date.today()}.docx",
            "Access-Control-Expose-Headers": "Content-Disposition",
        },
    )


def build_renter_bank_report_line(payment_info, service_name):
    output_line = ""
    line = "|{:10s}|{:8s}|{:9.5f}|{:8.6f}|{:8.6f}|{:9.2f}|{:7.2f}|{:9.2f}|{:9.2f}|"

    coefficient = payment_info["coefficient_value"]
    if payment_info["is_additional_coefficient_applied"]:
        coefficient = payment_info["additional_coefficient_value"]

    output_line += (
        line.format(
            "Отопление" if service_name == "heating" else "ГВС",
            f"{payment_info['payment_month']}.{payment_info['payment_year']}",
            payment_info["applied_rate_value"],
            payment_info[f"{service_name}_value"],
            coefficient,
            payment_info[f"{service_name}_cost"],
            payment_info[f"{service_name}_vat_cost"],
            payment_info[f"{service_name}_coefficient_value"],
            payment_info[f"{service_name}_total_cost"],
        )
        + "\n"
    )

    return output_line
